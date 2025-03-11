import sys
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
import tkinter as tk
from tkinter import messagebox


def get_template_path():
    """
    根据运行环境返回模板文件路径，
    若为 PyInstaller 打包后的环境，则从 sys._MEIPASS 中查找模板；
    否则，从项目的 templates 目录中查找模板文件。
    """
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
        template_path = os.path.join(base_path, "templates", "请假条模板.docx")
    else:
        base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
        template_path = os.path.join(base_dir, 'templates', '请假条模板.docx')
    return template_path


def create_leave_application(all_info):
    template_path = get_template_path()

    if not os.path.exists(template_path):
        print(f"模板文件不存在: {template_path}")
        return

    final_doc = Document()

    # 获取桌面路径并设置保存路径
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    save_path = os.path.join(desktop_path, "请假条.docx")

    if not os.access(desktop_path, os.W_OK):
        print(f"无法在桌面路径中写入: {desktop_path}")
        return

    for info in all_info:
        try:
            person_doc = Document(template_path)
            replacements = {
                "input1": info.get('学院', ''),
                "input2": info.get('专业', ''),
                "input3": info.get('班级', ''),
                "input4": info.get('姓名', ''),
                "input5": info.get('学号', ''),
                "input6": info.get('事由', ''),
                "input7": info.get('请假时间', ''),
                "y": str(info.get('日期（年）', '')),
                "m": str(info.get('日期（月）', '')),
                "d": str(info.get('日期（日）', ''))
            }

            # 替换模板中所有占位符
            for paragraph in person_doc.paragraphs:
                for run in paragraph.runs:
                    for placeholder, value in replacements.items():
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)

            # 将 person_doc 中的内容复制到最终文档中
            for paragraph in person_doc.paragraphs:
                p = final_doc.add_paragraph()
                p._element.clear_content()
                p._element.extend(paragraph._element)

                p.paragraph_format.line_spacing = Pt(18)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

                if any(phrase in p.text for phrase in ["您好！", "兹有", "望您原谅，予以批准。"]):
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                if "请假条" in p.text:
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        except Exception as e:
            print(f"处理假条时出错: {e}")
            continue

    try:
        final_doc.save(save_path)
        print(f"假条已成功保存到桌面，路径：{save_path}")
        show_file_saved_message(save_path)
    except Exception as e:
        print(f"保存文件时出错: {e}")


def show_file_saved_message(save_path):
    root = tk.Tk()
    root.withdraw()
    messagebox.showinfo("文件保存成功", f"假条已成功保存到：{save_path}")
    root.destroy()
